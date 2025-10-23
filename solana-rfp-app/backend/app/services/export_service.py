from typing import List, Dict, Any
from datetime import datetime
import io
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from app.models.schemas import QuestionResult, ProcessingResult


class ExportService:
    """Service for exporting Q&A results to PDF and DOCX formats"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_pdf_styles()
    
    def _setup_pdf_styles(self):
        """Setup custom PDF styles"""
        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor='#1a365d'
        ))
        
        # Question style
        self.styles.add(ParagraphStyle(
            name='Question',
            parent=self.styles['Normal'],
            fontSize=12,
            spaceAfter=6,
            spaceBefore=12,
            leftIndent=0,
            textColor='#2d3748',
            fontName='Helvetica-Bold'
        ))
        
        # Answer style
        self.styles.add(ParagraphStyle(
            name='Answer',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=20,
            leftIndent=20,
            textColor='#4a5568'
        ))
        
        # Empty answer style
        self.styles.add(ParagraphStyle(
            name='EmptyAnswer',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=20,
            leftIndent=20,
            textColor='#e53e3e',
            fontName='Helvetica-Oblique'
        ))
    
    def export_to_pdf(self, results: ProcessingResult, filename: str = None) -> bytes:
        """Export Q&A results to PDF format"""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"rfp_answers_{timestamp}.pdf"
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        # Build content
        story = []
        
        # Title
        title = Paragraph("RFP Questions & Answers", self.styles['CustomTitle'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Summary
        summary_text = f"""
        <b>Summary:</b><br/>
        Questions Processed: {results.questions_processed}<br/>
        Answers Found: {results.answers_found}<br/>
        Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}
        """
        summary = Paragraph(summary_text, self.styles['Normal'])
        story.append(summary)
        story.append(Spacer(1, 20))
        
        # Q&A pairs
        for i, result in enumerate(results.results, 1):
            # Question
            question_text = f"<b>Question {i}:</b> {result.question}"
            question = Paragraph(question_text, self.styles['Question'])
            story.append(question)
            
            # Answer
            if result.answer and result.answer.strip() and result.answer != "No answer found in knowledge base":
                answer_text = result.answer
                answer = Paragraph(answer_text, self.styles['Answer'])
            else:
                answer_text = "[Answer to be provided manually]"
                answer = Paragraph(answer_text, self.styles['EmptyAnswer'])
            
            story.append(answer)
            
            # Add spacing between Q&A pairs
            if i < len(results.results):
                story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_to_docx(self, results: ProcessingResult, filename: str = None) -> bytes:
        """Export Q&A results to DOCX format"""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"rfp_answers_{timestamp}.docx"
        
        doc = Document()
        
        # Title
        title = doc.add_heading('RFP Questions & Answers', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary
        doc.add_heading('Summary', level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run(f'Questions Processed: {results.questions_processed}\n')
        summary_para.add_run(f'Answers Found: {results.answers_found}\n')
        summary_para.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
        
        # Add spacing
        doc.add_paragraph()
        
        # Q&A pairs
        for i, result in enumerate(results.results, 1):
            # Question
            question_heading = doc.add_heading(f'Question {i}', level=2)
            question_para = doc.add_paragraph(result.question)
            
            # Answer
            answer_heading = doc.add_heading('Answer', level=3)
            if result.answer and result.answer.strip() and result.answer != "No answer found in knowledge base":
                answer_para = doc.add_paragraph(result.answer)
            else:
                answer_para = doc.add_paragraph()
                run = answer_para.add_run('[Answer to be provided manually]')
                run.italic = True
                run.font.color.rgb = None  # Red color for empty answers
            
            # Add spacing between Q&A pairs
            if i < len(results.results):
                doc.add_paragraph()
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def get_export_filename(self, format_type: str, custom_name: str = None) -> str:
        """Generate filename for export"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if custom_name:
            base_name = custom_name.replace(' ', '_').lower()
        else:
            base_name = "rfp_answers"
        
        return f"{base_name}_{timestamp}.{format_type.lower()}"
