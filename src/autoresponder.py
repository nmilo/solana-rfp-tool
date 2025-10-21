# src/autoresponder.py
import argparse, json, re, sys
from pathlib import Path
from datetime import datetime

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document

def normalize_text(s: str) -> str:
    s = s.lower()
    s = re.sub(r"[^\w\s\-\.]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s

# manje agresivni filteri (NE uklanjamo "as requested" itd.)
STOP_PREFIX = tuple([
    "from:", "sent:", "subject:", "to:", "cc:", "bcc:", "attachments:",
    "unsubscribe", "confidential", "disclaimer"
])

QUESTION_TRIGGERS = [
    "please provide", "please share", "kindly provide", "kindly share",
    "provide a", "provide the", "provide your",
    "please tell us", "let us know", "could you", "can you",
    "we would appreciate", "we’d appreciate", "we would like to request",
    "we'd like to request", "we request", "request if you can share",
    "we’d like to ask", "we would like to ask",
    "how does", "how do", "how is", "how are", "how will",
    "what is", "what are", "what’s", "which", "when", "why",
    "justify", "outline", "explain", "describe", "bsp:"
]

BULLET_RE = re.compile(r"^\s*(?:[\-\*\u2022]|\d+\)|\d+\.)\s+")
URL_ONLY_RE = re.compile(r"^https?://\S+$", re.I)

def clean_email_lines(raw: str):
    text = re.split(r"\nOn .+ wrote:\n", raw, flags=re.IGNORECASE)[0]
    lines = []
    for ln in text.splitlines():
        l = ln.rstrip()
        if not l: 
            continue
        low = l.strip().lower()
        if low.startswith(STOP_PREFIX):  # header/meta
            continue
        if URL_ONLY_RE.match(l.strip()):
            continue
        lines.append(l.strip())
    merged, buf = [], ""
    for l in lines:
        if BULLET_RE.match(l):
            if buf:
                merged.append(buf.strip()); buf = ""
            merged.append(l)
            continue
        if not buf:
            buf = l
        else:
            if not re.search(r"[\.!?]\s*$", buf):
                buf += " " + l
            else:
                merged.append(buf.strip()); buf = l
    if buf:
        merged.append(buf.strip())
    return merged

def looks_like_question(s: str) -> bool:
    low = s.lower().strip()
    if low.endswith("?"):
        return True
    if BULLET_RE.match(s) and any(tr in low for tr in QUESTION_TRIGGERS):
        return True
    if any(tr in low for tr in QUESTION_TRIGGERS):
        return True
    if re.match(r"^\s*(provide|outline|explain|describe|justify)\b", low):
        return True
    if re.match(r"^\s*(bsp|regulator)\s*:\s*", low):
        return True
    return False

def split_on_question_marks(s: str):
    parts = re.split(r"(\?)", s)
    out, acc = [], ""
    for seg in parts:
        acc += seg
        if seg == "?":
            out.append(acc.strip()); acc = ""
    if acc.strip():
        out.append(acc.strip())
    return out

def extract_questions(raw: str, debug=False):
    lines = clean_email_lines(raw)
    candidates = []
    for l in lines:
        for piece in split_on_question_marks(l):
            if looks_like_question(piece):
                candidates.append(piece)
    extra = []
    for l in lines:
        low = l.lower()
        if any(tr in low for tr in QUESTION_TRIGGERS) and len(low) > 25:
            extra.append(l)
    candidates += extra
    seen, qs = set(), []
    for q in candidates:
        qn = normalize_text(q)
        if len(qn) < 15: 
            continue
        if qn in seen:
            continue
        seen.add(qn)
        qs.append(q.strip())
    if debug:
        sys.stderr.write("\n[DEBUG] Detected questions:\n")
        for i, q in enumerate(qs, 1):
            sys.stderr.write(f"  {i}. {q}\n")
        sys.stderr.flush()
    return qs

def load_kb(path: Path):
    data = json.loads(path.read_text(encoding="utf-8"))
    for i, item in enumerate(data):
        item.setdefault("id", f"kb_{i}")
        item.setdefault("tags", [])
    return data

def build_retriever(kb_items):
    corpus = [normalize_text(x.get("question","") + " " + x.get("answer","") + " " + " ".join(x.get("tags", []))) for x in kb_items]
    vectorizer = TfidfVectorizer(ngram_range=(1,2), min_df=1)
    X = vectorizer.fit_transform(corpus)
    return vectorizer, X, corpus

def jaccard(a: set, b: set) -> float:
    if not a or not b: return 0.0
    inter = len(a & b)
    if inter == 0: return 0.0
    return inter / float(len(a | b))

def retrieve_answers(raw_email: str, kb_items, retriever, min_score=0.10, top_answers=1, debug=False):
    vectorizer, X, corpus = retriever
    questions = extract_questions(raw_email, debug=debug)
    results = []
    for q in questions:
        qn = normalize_text(q)
        qv = vectorizer.transform([qn])
        sims = cosine_similarity(qv, X)[0]
        q_tokens = set(qn.split())
        jac = [jaccard(q_tokens, set(doc.split())) for doc in corpus]
        fused = [(i, 0.7*float(sims[i]) + 0.3*float(jac[i])) for i in range(len(corpus))]
        fused.sort(key=lambda x: x[1], reverse=True)
        kept = [i for (i, s) in fused if s >= min_score]
        if not kept:
            results.append({"question": q, "answers": []})
            continue
        top = kept[:max(1, min(top_answers, 3))]
        answers = [kb_items[i]["answer"] for i in top]
        results.append({"question": q, "answers": answers})
    return results

def compose_email(counterparty: str, findings: list, sign_off: str = "Best regards,\nSolana Ecosystem Team"):
    now_utc = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%SZ")
    lines = []
    subj = f"Re: Your RFP Questions — {counterparty}" if counterparty else "Re: Your RFP Questions"
    lines.append(f"Subject: {subj}")
    lines.append("")
    greet = f"Hi {counterparty}," if counterparty else "Hello,"
    lines.append(greet)
    lines.append("")
    lines.append("Thanks for the detailed questions. Please find our consolidated responses below.")
    lines.append("")
    open_items = []
    for i, item in enumerate(findings, 1):
        lines.append(f"{i}) {item['question']}")
        if item["answers"]:
            for a in item["answers"]:
                lines.append(f"   • {a}")
        else:
            lines.append("   • [No canonical answer in KB]")
            open_items.append(item["question"])
        lines.append("")
    if open_items:
        lines.append("---")
        lines.append("Open items (add to KB and re-generate):")
        for q in open_items:
            lines.append(f"- {q}")
        lines.append("")
    lines.append("If helpful, we can share links to public documentation and audits, or schedule a technical walk-through.")
    lines.append("")
    lines.append(sign_off)
    lines.append("")
    lines.append(f"(Auto-generated preview, UTC {now_utc})")
    return "\n".join(lines)

def write_markdown(text: str, out_path: Path):
    out_path.write_text(text, encoding="utf-8")

def write_docx(text: str, out_path: Path):
    doc = Document()
    for line in text.splitlines():
        if line.startswith("Subject:"):
            p = doc.add_paragraph(); p.add_run(line).bold = True
        else:
            doc.add_paragraph(line)
    doc.save(out_path.as_posix())

def main():
    ap = argparse.ArgumentParser(description="RFP Auto-Responder")
    ap.add_argument("--kb", required=True)
    ap.add_argument("--in", dest="input_path")
    ap.add_argument("--out", required=True)
    ap.add_argument("--counterparty", default="")
    ap.add_argument("--min_score", type=float, default=0.10)
    ap.add_argument("--top_answers", type=int, default=1)
    ap.add_argument("--debug", action="store_true", help="Print detected questions to stderr")
    args = ap.parse_args()

    kb = load_kb(Path(args.kb))
    raw_email = Path(args.input_path).read_text(encoding="utf-8") if args.input_path else sys.stdin.read()

    retriever = build_retriever(kb)
    findings = retrieve_answers(
        raw_email, kb, retriever,
        min_score=args.min_score,
        top_answers=args.top_answers,
        debug=args.debug
    )

    body = compose_email(counterparty=args.counterparty, findings=findings)
    out_md = Path(args.out).with_suffix(".md")
    out_docx = Path(args.out).with_suffix(".docx")
    out_md.parent.mkdir(parents=True, exist_ok=True)
    write_markdown(body, out_md)
    write_docx(body, out_docx)

    print(json.dumps({
        "output_markdown": out_md.as_posix(),
        "output_docx": out_docx.as_posix(),
        "questions_detected": [f["question"] for f in findings]
    }, ensure_ascii=False))

if __name__ == "__main__":
    main()
